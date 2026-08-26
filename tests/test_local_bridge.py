from pathlib import Path

from docx import Document

from document_writer import validate_docx
from local_bridge import SessionStore, dispatch


def test_authentication_and_session_scope(tmp_path: Path) -> None:
    source = tmp_path / "source.txt"
    source.write_text("du lieu hoi thi", encoding="utf-8")
    sessions = SessionStore()

    bad = dispatch({"id": "1", "method": "agent.status", "token": "bad"}, token="secret", sessions=sessions)
    assert bad["ok"] is False
    assert bad["error"]["code"] == "AUTHENTICATION_FAILED"

    opened = dispatch({"id": "2", "method": "session.open", "token": "secret", "authorized_paths": [str(tmp_path)]}, token="secret", sessions=sessions)
    assert opened["ok"] is True
    session_id = opened["data"]["session_id"]

    denied = dispatch({"id": "3", "method": "file.metadata", "token": "secret", "session_id": session_id, "path": str(tmp_path.parent)}, token="secret", sessions=sessions)
    assert denied["ok"] is False
    assert denied["error"]["code"] == "AUTHORIZATION_REQUIRED"


def test_source_overwrite_and_structured_docx(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    doc = Document()
    doc.add_heading("Kế hoạch thử nghiệm", level=1)
    doc.add_paragraph("Dữ liệu người dùng")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Chỉ tiêu"
    table.cell(0, 1).text = "Giá trị"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "10"
    doc.save(source)
    original = source.read_bytes()
    sessions = SessionStore()
    opened = dispatch({"id": "1", "method": "session.open", "token": "secret", "authorized_paths": [str(tmp_path)]}, token="secret", sessions=sessions)
    sid = opened["data"]["session_id"]

    read = dispatch({"id": "2", "method": "document.read", "token": "secret", "session_id": sid, "path": str(source)}, token="secret", sessions=sessions)
    assert read["ok"] is True
    assert len(read["data"]["tables"]) == 1
    assert any(item["type"] == "table" for item in read["data"]["body"])

    blocked = dispatch({"id": "3", "method": "document.write", "token": "secret", "session_id": sid, "path": str(source), "source_path": str(source), "document": read["data"]}, token="secret", sessions=sessions)
    assert blocked["ok"] is False
    assert blocked["error"]["code"] == "SOURCE_OVERWRITE_BLOCKED"
    assert source.read_bytes() == original

    output = tmp_path / "rewritten.docx"
    written = dispatch({"id": "4", "method": "document.write", "token": "secret", "session_id": sid, "path": str(output), "source_path": str(source), "document": read["data"]}, token="secret", sessions=sessions)
    assert written["ok"] is True
    assert written["data"]["validation"]["valid"] is True
    assert validate_docx(output)["valid"] is True
    assert output.exists()
    assert source.read_bytes() == original


def test_protocol_error_does_not_leak_exception(tmp_path: Path) -> None:
    sessions = SessionStore()
    response = dispatch({"id": "1", "method": "file.metadata", "token": "secret", "session_id": "missing", "path": str(tmp_path)}, token="secret", sessions=sessions)
    assert response["ok"] is False
    assert response["error"]["code"] == "SESSION_NOT_FOUND"
    assert "Traceback" not in str(response)
