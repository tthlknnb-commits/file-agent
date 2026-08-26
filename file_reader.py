from __future__ import annotations

import hashlib
import json
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".xml", ".html", ".htm"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | {".docx", ".xlsx", ".pdf"}


@dataclass
class Evidence:
    source_uri: str
    title: str
    file_type: str
    locator: str | None
    content: str
    sha256: str
    status: str = "UNVERIFIED"
    metadata: dict[str, Any] = field(default_factory=dict)

    def as_dict(self) -> dict[str, Any]:
        return asdict(self)


def _hash(content: str) -> str:
    return hashlib.sha256(content.encode("utf-8", errors="replace")).hexdigest()


def _read_text(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1258", "cp1252"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return ""


def read_file(path: str | Path) -> list[Evidence]:
    """Extract evidence without changing the source file."""
    p = Path(path).expanduser()
    suffix = p.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        return []
    if not p.is_file():
        return []

    chunks: list[tuple[str | None, str]] = []
    if suffix in TEXT_EXTENSIONS:
        text = _read_text(p)
        if text:
            chunks.append((None, text))
    elif suffix == ".docx":
        try:
            from docx import Document
            doc = Document(p)
            text = "\n".join(x.text for x in doc.paragraphs if x.text.strip())
            for table_index, table in enumerate(doc.tables, 1):
                rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
                text += f"\n[TABLE {table_index}]\n" + "\n".join(rows)
            if text.strip():
                chunks.append(("document", text))
        except Exception as exc:
            chunks.append(("read_error", f"READ_ERROR: {exc}"))
    elif suffix == ".xlsx":
        try:
            from openpyxl import load_workbook
            workbook = load_workbook(p, data_only=False, read_only=True)
            for sheet in workbook.worksheets:
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    rows.append(" | ".join("" if value is None else str(value) for value in row))
                chunks.append((f"sheet:{sheet.title}", "\n".join(rows)))
        except Exception as exc:
            chunks.append(("read_error", f"READ_ERROR: {exc}"))
    elif suffix == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(p))
            for page_no, page in enumerate(reader.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    chunks.append((f"page:{page_no}", text))
        except Exception as exc:
            chunks.append(("read_error", f"READ_ERROR: {exc}"))

    return [
        Evidence(str(p), p.name, suffix.lstrip("."), locator, content, _hash(content),
                 "USER_SUPPLIED", {"size": p.stat().st_size})
        for locator, content in chunks if content
    ]


def read_paths(paths: list[str], recursive: bool = True) -> list[Evidence]:
    evidence: list[Evidence] = []
    for raw in paths:
        p = Path(raw).expanduser()
        if not p.exists():
            continue
        candidates = p.rglob("*") if p.is_dir() and recursive else [p]
        for candidate in candidates:
            if candidate.is_file() and candidate.suffix.lower() in SUPPORTED_EXTENSIONS:
                evidence.extend(read_file(candidate))
    return evidence


def evidence_ledger(evidence: list[Evidence]) -> list[dict[str, Any]]:
    return [item.as_dict() for item in evidence]


def save_ledger(evidence: list[Evidence], output: str | Path) -> None:
    Path(output).write_text(json.dumps(evidence_ledger(evidence), ensure_ascii=False, indent=2), encoding="utf-8")
