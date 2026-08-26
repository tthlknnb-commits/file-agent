from __future__ import annotations

from pathlib import Path
from typing import Any


def write_text(content: str, output: str | Path) -> Path:
    path = Path(output)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    return path


def read_docx(source: str | Path) -> dict[str, Any]:
    from docx import Document
    path = Path(source)
    document = Document(path)
    body: list[dict[str, Any]] = []
    for child in document.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            paragraph = next((p for p in document.paragraphs if p._p is child), None)
            if paragraph is not None:
                body.append({"type": "paragraph", "text": paragraph.text, "style": paragraph.style.name})
        elif tag == "tbl":
            table = next((t for t in document.tables if t._tbl is child), None)
            if table is not None:
                body.append({"type": "table", "rows": [[cell.text for cell in row.cells] for row in table.rows]})
    sections = []
    for section in document.sections:
        sections.append({"top_margin": section.top_margin.inches, "bottom_margin": section.bottom_margin.inches, "left_margin": section.left_margin.inches, "right_margin": section.right_margin.inches, "header": [p.text for p in section.header.paragraphs], "footer": [p.text for p in section.footer.paragraphs]})
    props = document.core_properties
    return {"source": str(path.resolve()), "body": body, "paragraphs": [p.text for p in document.paragraphs], "tables": [{"rows": [[cell.text for cell in row.cells] for row in table.rows]} for table in document.tables], "sections": sections, "core_properties": {"title": props.title, "subject": props.subject, "author": props.author, "keywords": props.keywords, "comments": props.comments}}


def write_docx(document_or_content: str | dict[str, Any], output: str | Path, title: str | None = None, *, overwrite: bool = False) -> Path:
    from docx import Document
    from docx.shared import Pt
    path = Path(output)
    if path.exists() and not overwrite:
        raise FileExistsError(f"output exists: {path}")
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    if isinstance(document_or_content, str):
        if title:
            document.add_heading(title, level=1)
        for block in document_or_content.split("\n\n"):
            if block.strip():
                document.add_paragraph(block.strip())
    else:
        data = document_or_content
        if title and not data.get("title"):
            document.add_heading(title, level=1)
        for item in data.get("body", []):
            if item.get("type") == "table":
                rows = item.get("rows", [])
                if rows:
                    table = document.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                    for i, row in enumerate(rows):
                        for j, value in enumerate(row):
                            table.cell(i, j).text = str(value)
            else:
                p = document.add_paragraph(str(item.get("text", "")))
                if item.get("style"):
                    try:
                        p.style = item["style"]
                    except (KeyError, ValueError):
                        pass
        if data.get("sections"):
            first = data["sections"][0]
            section = document.sections[0]
            for key in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
                if first.get(key) is not None:
                    from docx.shared import Inches
                    setattr(section, key, Inches(float(first[key])))
            for text in first.get("header", []):
                section.header.add_paragraph(str(text))
            for text in first.get("footer", []):
                section.footer.add_paragraph(str(text))
        props = document.core_properties
        for key in ("title", "subject", "author", "keywords", "comments"):
            value = data.get("core_properties", {}).get(key)
            if value is not None:
                setattr(props, key, value)
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)
    document.save(path)
    return path


def validate_docx(source: str | Path, *, expected_text: str | None = None) -> dict[str, Any]:
    try:
        data = read_docx(source)
        text = "\n".join(data["paragraphs"])
        if expected_text is not None and expected_text not in text:
            return {"valid": False, "reason": "expected text missing", "paragraph_count": len(data["paragraphs"]), "table_count": len(data["tables"])}
        return {"valid": True, "paragraph_count": len(data["paragraphs"]), "table_count": len(data["tables"]), "section_count": len(data["sections"]), "path": str(Path(source).resolve())}
    except Exception:
        return {"valid": False, "reason": "DOCX could not be opened"}


def write_csv(rows: list[dict[str, Any]], output: str | Path) -> Path:
    import csv
    path = Path(output)
    path.parent.mkdir(parents=True, exist_ok=True)
    if not rows:
        path.write_text("", encoding="utf-8-sig")
        return path
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)
    return path
